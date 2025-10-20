from docx import Document
import json


doc = Document('testPR.docx')

doc_formated = Document()


doc_tables = doc.tables

print(len(doc_tables))

table_template = {
    "rows" : ,
    "columns" : ,

}

to_json = {
    "table1" : doc_tables[0],
}

with open('sw_templates.json', 'w') as f:
    f.write(json.dumps(to_json))



for n in range(len(doc_tables)):
    table = doc_tables[n]
    doc_formated.add_table(rows=len(table.rows), cols=len(table.columns))
    for i, row in enumerate(table.rows):
        for j, column in enumerate(table.columns):
            print(f"{table.cell(i, j).text}, ( table: {n + 1}, row: {i}, column: {j})")
            doc_formated.tables[n].cell(i, j).add_paragraph(text=table.cell(i, j).text)

"""
for i in range(len(doc.paragraphs)):
    text = doc.paragraphs[i].text
    # print(text)
    doc_formated.add_paragraph(text, style="Normal")



doc_formated.save("testPRFormated.docx")
"""